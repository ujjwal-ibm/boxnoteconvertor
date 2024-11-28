from bs4 import BeautifulSoup
from docx import Document
from docx.shared import RGBColor, Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import docx.table
import re
from html.parser import HTMLParser
from pathlib import Path
import os
from typing import Dict, Any, Optional

from boxnotetodocx.utils.logger import get_logger

logger = get_logger(__name__)

# Configuration constants
INDENT = 0.25
LIST_INDENT = 0.5
BLOCKQUOTE_INDENT = 0.5
MAX_INDENT = 5.5
DEFAULT_TABLE_STYLE = 'TableGrid'
DEFAULT_PARAGRAPH_STYLE = None

# Style mappings
font_styles = {
    'b': 'bold',
    'strong': 'bold',
    'em': 'italic',
    'i': 'italic',
    'u': 'underline',
    's': 'strike',
    'sup': 'superscript',
    'sub': 'subscript',
    'th': 'bold',
    'code': 'bold',
}

font_names = {
    'code': 'Courier New',
    'pre': 'Courier New',
}
styles = {
    'LIST_BULLET': 'List Bullet',
    'LIST_NUMBER': 'List Number',
}


def delete_paragraph(paragraph):
    """Safely delete a paragraph"""
    try:
        p = paragraph._element
        parent = p.getparent()
        if parent is not None:
            parent.remove(p)
        paragraph._p = paragraph._element = None
    except Exception as e:
        logger.error(f"Error deleting paragraph: {e}")

class HtmlToDocx(HTMLParser):
    def __init__(self, workdir=None):
        super().__init__()
        self.workdir = workdir
        self.table_row_selectors = [
            'table > tr',
            'table > thead > tr',
            'table > tbody > tr',
            'table > tfoot > tr'
        ]
        self.table_style = 'TableGrid'
        self.paragraph_style = None
        self.reset_state()

    def reset_state(self):
        """Reset parser state"""
        self.tags = {
            'span': [],
            'list': [],
            'blockquote': [],
            'table': []
        }
        self.doc = None
        self.paragraph = None
        self.run = None
        self.skip = False
        self.skip_tag = None
        self.instances_to_skip = 0
        self.blockquote = False
        self.style = False
        self.current_table = None
        self.in_header = False
        self.current_tag_content = []  # To store content between tags
        self.current_list_type = None
        self.current_list_level = 0
        self.current_table_content = []
        self.current_table_row = None
        self.current_table_cell = None
        HTMLParser.reset(self)
        

    def parse_style_string(self, style_string: str) -> dict:
        """Parse HTML style string to dictionary"""
        try:
            if not style_string:
                return {}
            
            styles = {}
            for style in style_string.split(';'):
                if ':' not in style:
                    continue
                    
                key, value = style.split(':', 1)
                key = key.strip().lower()
                value = value.strip()
                
                if 'rgb' in value:
                    rgb = re.findall(r'\d+', value)
                    if len(rgb) == 3:
                        value = '#{:02x}{:02x}{:02x}'.format(*map(int, rgb))
                        
                styles[key] = value
                
            return styles
            
        except Exception as e:
            logger.warning(f"Error parsing style string: {e}")
            return {}

    def handle_starttag(self, tag, attrs):
        """Handle opening HTML tags"""
        if self.skip:
            return

        # Skip style tags completely
        if tag == 'style':
            self.skip = True
            self.skip_tag = tag
            return

        try:
            attrs_dict = dict(attrs)
            
            if tag == 'p':
                self.paragraph = self.doc.add_paragraph()
                if 'style' in attrs_dict:
                    style = self.parse_style_string(attrs_dict['style'])
                    self.apply_paragraph_style(style)
                    
            elif tag in ['ul', 'ol']:
                self.current_list_type = tag
                self.current_list_level += 1
                
            elif tag == 'li':
                self.handle_list_item()
                
            elif tag == 'table':
                self.current_table_content = []
                self.in_header = False
            elif tag == 'tr':
                self.current_table_row = []
            elif tag in ['td', 'th']:
                self.in_header = tag == 'th'
                self.current_table_cell = []
                
            elif tag in ['strong', 'b', 'em', 'i', 'u', 's']:
                if not self.paragraph:
                    self.paragraph = self.doc.add_paragraph()
                self.run = self.paragraph.add_run()
                self.apply_font_style(tag)
                
            elif tag == 'a' and 'href' in attrs_dict:
                if not self.paragraph:
                    self.paragraph = self.doc.add_paragraph()
                self.handle_link(attrs_dict['href'])
                
            elif tag == 'img' and 'src' in attrs_dict:
                self.handle_image(attrs_dict)
                
        except Exception as e:
            logger.error(f"Error handling start tag {tag}: {e}")

    def handle_endtag(self, tag):
        """Handle closing HTML tags"""
        if self.skip:
            if tag == self.skip_tag:
                self.skip = False
                self.skip_tag = None
            return

        try:
            if tag in ['ul', 'ol']:
                self.current_list_level -= 1
                if self.current_list_level == 0:
                    self.current_list_type = None
                    
            elif tag == 'table':
                self.process_table()
                self.current_table_content = []
            elif tag == 'tr':
                if self.current_table_row:
                    self.current_table_content.append(self.current_table_row)
                self.current_table_row = None
            elif tag in ['td', 'th']:
                if self.current_table_cell and self.current_table_row is not None:
                    cell_content = ''.join(self.current_table_cell)
                    self.current_table_row.append((cell_content, self.in_header))
                self.current_table_cell = None
                
        except Exception as e:
            logger.error(f"Error handling end tag {tag}: {e}")

    def handle_data(self, data):
        """Handle text content"""
        if self.skip or not data.strip():
            return

        try:
            if not self.paragraph:
                self.paragraph = self.doc.add_paragraph()
                
            if not self.run:
                self.run = self.paragraph.add_run(data)
            else:
                self.run.add_text(data)
                
        except Exception as e:
            logger.error(f"Error handling data: {e}")

    def parse_html_file(self, input_file: str, output_file: str = None) -> None:
        """Parse HTML file to DOCX with enhanced error handling"""
        try:
            # Read HTML file
            with open(input_file, 'r', encoding='utf-8') as infile:
                html = infile.read()

            # Initialize document and state
            self.reset_state()
            self.doc = Document()

            # Parse HTML
            self.feed(html)

            # Clean up empty paragraphs
            for paragraph in self.doc.paragraphs[:]:
                if not paragraph.text and not paragraph._element.xpath('.//w:drawing'):
                    self.delete_paragraph(paragraph)

            # Save document
            if not output_file:
                path, filename = os.path.split(input_file)
                output_file = f'{path}/new_docx_file_{filename}'

            # Ensure no double extension
            output_file = output_file.replace('.docx', '')
            self.doc.save(f'{output_file}.docx')

        except Exception as e:
            logger.error(f"Error parsing HTML file: {e}")
            raise


    def apply_paragraph_style(self, style: dict):
        """Apply style to paragraph"""
        try:
            if 'text-align' in style:
                align = style['text-align']
                if align == 'center':
                    self.paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                elif align == 'right':
                    self.paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                elif align == 'justify':
                    self.paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    
            if 'margin-left' in style:
                margin = style['margin-left']
                value = float(re.findall(r'[\d.]+', margin)[0])
                if 'px' in margin:
                    value = value / 96  # Convert pixels to inches
                elif 'pt' in margin:
                    value = value / 72  # Convert points to inches
                self.paragraph.paragraph_format.left_indent = Inches(min(value, MAX_INDENT))
                
        except Exception as e:
            logger.error(f"Error applying paragraph style: {e}")

    def apply_font_style(self, tag: str):
        """Apply font style to run"""
        try:
            if tag in font_styles:
                setattr(self.run.font, font_styles[tag], True)
            if tag in font_names:
                self.run.font.name = font_names[tag]
        except Exception as e:
            logger.error(f"Error applying font style: {e}")

    def handle_list_item(self):
        """Handle list item creation"""
        try:
            style = styles['LIST_BULLET'] if self.current_list_type == 'ul' else styles['LIST_NUMBER']
            self.paragraph = self.doc.add_paragraph(style=style)
            self.paragraph.paragraph_format.left_indent = Inches(min(self.current_list_level * LIST_INDENT, MAX_INDENT))
            self.run = None
        except Exception as e:
            logger.error(f"Error handling list item: {e}")


    def apply_text_style(self, run, style_dict: dict):
        """Apply text styling with enhanced color support"""
        try:
            if 'color' in style_dict:
                color = style_dict['color']
                if color.startswith('#'):
                    # Convert hex to RGB
                    rgb = tuple(int(color[i:i+2], 16) for i in (1, 3, 5))
                    run.font.color.rgb = RGBColor(*rgb)
                    
            if 'background-color' in style_dict:
                color = style_dict['background-color']
                if color.startswith('#'):
                    rgb = tuple(int(color[i:i+2], 16) for i in (1, 3, 5))
                    shd = OxmlElement('w:shd')
                    shd.set(qn('w:fill'), '{:02x}{:02x}{:02x}'.format(*rgb))
                    run._element.rPr.append(shd)
                    
            if 'font-size' in style_dict:
                size = style_dict['font-size']
                if 'pt' in size:
                    pt_size = float(size.replace('pt', ''))
                    run.font.size = Pt(pt_size)
                elif 'px' in size:
                    px_size = float(size.replace('px', ''))
                    run.font.size = Pt(px_size * 0.75)
                    
        except Exception as e:
            logger.warning(f"Error applying text style: {e}")

    def handle_table(self, table_soup) -> None:
        """Handle table conversion with proper error handling"""
        try:
            if table_soup is None:
                logger.warning("Empty table found, skipping")
                return

            rows, cols = self.get_table_dimensions(table_soup)
            if rows == 0 or cols == 0:
                logger.warning("Table with no dimensions found, skipping")
                return

            # Create table
            table = self.doc.add_table(rows, cols)
            if self.table_style:
                try:
                    table.style = self.table_style
                except KeyError:
                    logger.warning(f"Table style '{self.table_style}' not found")

            # Process rows
            current_row = 0
            for row_elem in self.get_table_rows(table_soup):
                if current_row >= rows:
                    break

                # Process cells
                cells = row_elem.find_all(['th', 'td'], recursive=False)
                current_col = 0

                for cell_elem in cells:
                    if current_col >= cols:
                        break

                    try:
                        cell = table.cell(current_row, current_col)

                        # Handle colspan and rowspan
                        colspan = int(cell_elem.get('colspan', 1))
                        rowspan = int(cell_elem.get('rowspan', 1))

                        if colspan > 1 or rowspan > 1:
                            cell.merge(
                                table.cell(
                                    current_row + rowspan - 1,
                                    current_col + colspan - 1
                                )
                            )

                        # Process cell content
                        cell_html = str(cell_elem)
                        self.add_html_to_cell(cell_html, cell)

                        current_col += colspan

                    except Exception as e:
                        logger.warning(f"Error processing table cell: {e}")
                        current_col += 1

                current_row += 1

        except Exception as e:
            logger.error(f"Error handling table: {e}")
            self.paragraph = self.doc.add_paragraph()


    def handle_list(self, list_type: str, content):
        """Enhanced list handling with proper indentation"""
        try:
            self.tags['list'].append(list_type)
            list_depth = len(self.tags['list'])
            
            if list_type == 'ol':
                style = 'List Number'
            else:
                style = 'List Bullet'
                
            self.paragraph = self.doc.add_paragraph(style=style)
            self.paragraph.paragraph_format.left_indent = Inches(min(list_depth * LIST_INDENT, MAX_INDENT))
            self.paragraph.paragraph_format.line_spacing = 1
            
            # Process list content
            self.process_content(content)
            
            self.tags['list'].pop()
            
        except Exception as e:
            logger.warning(f"Error handling list: {e}")

    def handle_image(self, img_elem):
        """Enhanced image handling with size support"""
        try:
            src = img_elem.get('src', '')
            if not src:
                return
                
            image_path = self.workdir / Path(src) if self.workdir else Path(src)
            
            if not image_path.exists():
                logger.warning(f"Image not found: {image_path}")
                return
                
            width = img_elem.get('width')
            height = img_elem.get('height')
            
            if width and height:
                self.doc.add_picture(str(image_path), width=Inches(float(width)/96), height=Inches(float(height)/96))
            else:
                self.doc.add_picture(str(image_path))
                
        except Exception as e:
            logger.warning(f"Error handling image: {e}")

    def handle_link(self, href: str, text: str):
        """Enhanced link handling with better formatting"""
        try:
            if not href or not text:
                return
                
            # Create the hyperlink element
            rel_id = self.paragraph.part.relate_to(
                href,
                docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK,
                is_external=True
            )

            hyperlink = OxmlElement('w:hyperlink')
            hyperlink.set(qn('r:id'), rel_id)

            # Create the text run
            run = self.paragraph.add_run()
            rPr = OxmlElement('w:rPr')

            # Add color (blue)
            color = OxmlElement('w:color')
            color.set(qn('w:val'), '0000EE')
            rPr.append(color)

            # Add underline
            underline = OxmlElement('w:u')
            underline.set(qn('w:val'), 'single')
            rPr.append(underline)

            run._r.append(rPr)
            run._r.text = text

            hyperlink.append(run._r)
            self.paragraph._p.append(hyperlink)

        except Exception as e:
            logger.warning(f"Error handling link: {e}")
            # Fallback to plain text
            self.paragraph.add_run(text)

    def handle_code_block(self, content: str):
        """Enhanced code block handling with proper formatting"""
        try:
            self.paragraph = self.doc.add_paragraph()
            run = self.paragraph.add_run(content)
            run.font.name = 'Courier New'
            run.font.size = Pt(10)
            
            # Add gray background
            shd = OxmlElement('w:shd')
            shd.set(qn('w:fill'), 'F0F0F0')
            run._element.rPr.append(shd)
            
            # Add border
            pPr = self.paragraph._p.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            
            borders = ['top', 'left', 'bottom', 'right']
            for border in borders:
                borderElement = OxmlElement(f'w:{border}')
                borderElement.set(qn('w:val'), 'single')
                borderElement.set(qn('w:sz'), '4')
                borderElement.set(qn('w:space'), '0')
                borderElement.set(qn('w:color'), 'auto')
                pBdr.append(borderElement)
                
            pPr.append(pBdr)
            
        except Exception as e:
            logger.warning(f"Error handling code block: {e}")

    def handle_blockquote(self, content):
        """Enhanced blockquote handling with proper styling"""
        try:
            self.blockquote = True
            self.paragraph = self.doc.add_paragraph()
            self.paragraph.paragraph_format.left_indent = Inches(BLOCKQUOTE_INDENT)
            self.paragraph.paragraph_format.right_indent = Inches(BLOCKQUOTE_INDENT)
            self.paragraph.paragraph_format.line_spacing = 1
            
            # Add left border
            pPr = self.paragraph._p.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            left = OxmlElement('w:left')
            left.set(qn('w:val'), 'single')
            left.set(qn('w:sz'), '24')
            left.set(qn('w:space'), '0')
            left.set(qn('w:color'), 'CCCCCC')
            pBdr.append(left)
            pPr.append(pBdr)
            
            self.process_content(content)
            self.blockquote = False
            
        except Exception as e:
            logger.warning(f"Error handling blockquote: {e}")

    def parse_html_file(self, input_file: str, output_file: str = None):
        """Parse HTML file to DOCX with enhanced error handling"""
        try:
            with open(input_file, 'r', encoding='utf-8') as infile:
                html = infile.read()
                
            self.reset_state()
            self.doc = Document()
            
            # Parse and process HTML
            self.soup = BeautifulSoup(html, 'html.parser')
            self.tables = self.get_tables()
            self.table_no = 0
            self.feed(str(self.soup))
            
            # Remove empty paragraphs
            for paragraph in self.doc.paragraphs:
                if not paragraph.text and not paragraph._element.xpath('.//w:drawing'):
                    self.delete_paragraph(paragraph)
                    
            # Save document
            if not output_file:
                path, filename = os.path.split(input_file)
                output_file = f'{path}/new_docx_file_{filename}'
                
            self.doc.save(f'{output_file}.docx')
            
        except Exception as e:
            logger.error(f"Error parsing HTML file: {e}")
            raise

    def delete_paragraph(self, paragraph):
        """Safely delete a paragraph"""
        try:
            p = paragraph._element
            parent = p.getparent()
            if parent is not None:
                parent.remove(p)
            paragraph._p = paragraph._element = None
        except Exception as e:
            logger.warning(f"Error deleting paragraph: {e}")

    def get_tables(self):
        """Get all tables excluding nested ones"""
        try:
            all_tables = self.soup.find_all('table')
            top_level_tables = []
            
            for table in all_tables:
                if not table.find_parent('table'):
                    top_level_tables.append(table)
                    
            return top_level_tables
            
        except Exception as e:
            logger.warning(f"Error getting tables: {e}")
            return []

    def get_table_dimensions(self, table_soup) -> tuple[int, int]:
        """Get table dimensions with error handling"""
        try:
            rows = self.get_table_rows(table_soup)
            max_cols = 0

            for row in rows:
                cols = self.get_table_columns(row)
                row_width = sum(int(col.get('colspan', 1)) for col in cols)
                max_cols = max(max_cols, row_width)

            return len(rows), max_cols

        except Exception as e:
            logger.warning(f"Error getting table dimensions: {e}")
            return 0, 0

    def add_html_to_cell(self, html: str, cell):
        """Add HTML content to table cell with proper formatting"""
        try:
            if not isinstance(cell, docx.table._Cell):
                raise ValueError('Second argument must be a table cell')
                
            # Remove default paragraph if empty
            if cell.paragraphs and not cell.paragraphs[0].text:
                self.delete_paragraph(cell.paragraphs[0])
                
            self.doc = cell
            soup = BeautifulSoup(html, 'html.parser')
            self.feed(str(soup))
            
            # Ensure cell has at least one paragraph
            if not cell.paragraphs:
                cell.add_paragraph()
                
        except Exception as e:
            logger.error(f"Error adding HTML to cell: {e}")
            if not cell.paragraphs:
                cell.add_paragraph()

    def process_table(self):
        """Process collected table content"""
        try:
            if not self.current_table_content:
                return

            # Calculate dimensions
            rows = len(self.current_table_content)
            cols = max(len(row) for row in self.current_table_content) if rows > 0 else 0

            if rows == 0 or cols == 0:
                return

            # Create table
            table = self.doc.add_table(rows, cols)
            if self.table_style:
                try:
                    table.style = self.table_style
                except KeyError:
                    logger.warning(f"Table style '{self.table_style}' not found")

            # Fill table
            for row_idx, row in enumerate(self.current_table_content):
                for col_idx, (content, is_header) in enumerate(row):
                    if col_idx < cols:  # Ensure we don't exceed table dimensions
                        cell = table.cell(row_idx, col_idx)
                        cell.text = content
                        if is_header:
                            run = cell.paragraphs[0].runs[0]
                            run.bold = True

        except Exception as e:
            logger.error(f"Error processing table: {e}")
