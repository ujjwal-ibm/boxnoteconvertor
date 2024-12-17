"""Handler for converting HTML content to DOCX format."""
from typing import Dict, Optional, Union
from pathlib import Path
import re
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
from docx.text.run import Run
from PIL import Image
from ..utils.logger import setup_logger
from ..utils.constants import (
    DEFAULT_TABLE_STYLE,
    DEFAULT_IMAGE_WIDTH,
    DEFAULT_FONT_SIZE,
    DEFAULT_FONT_NAME
)

logger = setup_logger(__name__)

class DOCXHandler:
    """Handles conversion of HTML content to DOCX format."""
    
    def __init__(self) -> None:
        """Initialize DOCX handler."""
        self.document = Document()
        self.output_dir: Optional[Path] = None
        self.assets_dir: Optional[Path] = None
        self._setup_document()
    
    def _setup_document(self) -> None:
        """Set up default document styles."""
        style = self.document.styles['Normal']
        font = style.font
        font.name = DEFAULT_FONT_NAME
        font.size = Pt(DEFAULT_FONT_SIZE)
    
    def convert_html_to_docx(self, html_content: str, output_path: Union[str, Path], assets_dir: Union[str, Path] = None) -> Path:
        """
        Convert HTML content to DOCX format.
        
        Args:
            html_content: HTML string to convert
            output_path: Path for output DOCX file
            assets_dir: Directory containing images and other assets

        Returns:
            Path to created DOCX file
            
        Raises:
            ValueError: If HTML content is invalid
        """
        try:
            output_path = Path(output_path)
            self.assets_dir = Path(assets_dir) if assets_dir else output_path.parent
            soup = BeautifulSoup(html_content, 'html.parser')
            
            if not soup.body:
                raise ValueError("Invalid HTML: no body tag found")
                
            self._process_elements(soup.body)
            
            output_path.parent.mkdir(parents=True, exist_ok=True)
            self.document.save(str(output_path))
            
            logger.info(f"Created DOCX file: {output_path}")
            return output_path
            
        except Exception as e:
            logger.error(f"DOCX conversion failed: {str(e)}")
            raise
    
    def _process_elements(self, parent: BeautifulSoup) -> None:
        """Process HTML elements recursively."""
        for element in parent.children:
            if not hasattr(element, 'name'):
                continue
                
            # Add logging to track element processing
            if element.name:
                logger.debug(f"Processing element: {element.name}")  # Add this line
                
            handler = self._element_handlers.get(element.name)
            if handler:
                handler(self, element)
                if element.name == 'img':  # Add this line
                    logger.info(f"Processed image element: {element}")  # Add this line
            else:
                logger.debug(f"Unhandled element type: {element.name}")
                
    def _handle_paragraph(self, element: BeautifulSoup) -> None:
        """Handle paragraph element."""
        p = self.document.add_paragraph()
        self._apply_styles(element, p)
        self._process_inline_elements(element, p)
    
    def _handle_heading(self, element: BeautifulSoup) -> None:
        """Handle heading element."""
        level = int(element.name[1])  # h1 -> 1, h2 -> 2, etc.
        self.document.add_heading(element.get_text(), level=level)
    
    def _handle_list(self, element: BeautifulSoup, ordered: bool = False) -> None:
        """Handle list element."""
        style = 'List Number' if ordered else 'List Bullet'
        for item in element.find_all('li', recursive=False):
            p = self.document.add_paragraph(style=style)
            self._process_inline_elements(item, p)
    
    def _handle_table(self, element: BeautifulSoup) -> None:
        """Handle table element."""
        rows = element.find_all('tr', recursive=False)
        if not rows:
            return
            
        cols = max(len(row.find_all(['td', 'th'], recursive=False)) for row in rows)
        table = self.document.add_table(rows=len(rows), cols=cols)
        table.style = DEFAULT_TABLE_STYLE
        
        for i, row in enumerate(rows):
            cells = row.find_all(['td', 'th'], recursive=False)
            for j, cell in enumerate(cells):
                table_cell = table.cell(i, j)
                self._process_inline_elements(cell, table_cell.paragraphs[0])
    
    def _handle_image(self, element: BeautifulSoup) -> None:
        """Handle image element."""
        try:
            logger.info("Image handler called")  # Add this line
            src = element.get('src')
            logger.info(f"Image src: {src}")  # Add this line
            
            if not src:
                logger.warning("No src attribute found in image element")
                return

            img_path = self.assets_dir / src
            logger.info(f"Looking for image at: {img_path}")  # Add this line
            
            if img_path.exists():
                logger.info(f"Found image file at: {img_path}")  # Add this line
                # Add image to document
                self.document.add_picture(str(img_path), width=Inches(6))
                self.document.add_paragraph()  # Add spacing after image
                logger.info("Successfully added image to document")  # Add this line
            else:
                logger.error(f"Image file not found: {img_path}")

        except Exception as e:
            logger.error(f"Error in image handling: {str(e)}")
            import traceback
            logger.error(traceback.format_exc())  # Add full traceback
    
    def _process_inline_elements(
        self,
        element: BeautifulSoup,
        paragraph: Paragraph
    ) -> None:
        """Process inline elements within a paragraph."""
        for child in element.children:
            if not hasattr(child, 'name') or child.name is None:
                # Handle pure text nodes
                if str(child).strip():  # Only add non-empty text
                    run = paragraph.add_run(str(child))
                continue
            
            # Add this block to handle images inside paragraphs
            if child.name == 'img':
                logger.info("Found image in paragraph")
                self._handle_image(child)
                continue
                
            # Handle other element nodes
            if child.name == 'a':
                self._add_hyperlink(paragraph, child.get('href', '#'), child.get_text())
            elif child.name in ['strong', 'b']:
                run = paragraph.add_run(child.get_text())
                run.bold = True
            elif child.name in ['em', 'i']:
                run = paragraph.add_run(child.get_text())
                run.italic = True
            elif child.name == 'u':
                run = paragraph.add_run(child.get_text())
                run.underline = True
            elif child.name == 'br':
                paragraph.add_run().add_break()
            else:
                run = paragraph.add_run(child.get_text())
            
            # Only apply styles to element nodes
            if child.name:
                self._apply_styles(child, run if 'run' in locals() else None)
    
    def _add_hyperlink(self, paragraph: Paragraph, url: str, text: str) -> None:
        """Add hyperlink to paragraph with proper formatting."""
        try:
            part = paragraph.part
            r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
            
            # Create hyperlink element
            hyperlink = OxmlElement('w:hyperlink')
            hyperlink.set(qn('r:id'), r_id)
            
            # Create run element
            run = OxmlElement('w:r')
            
            # Add style to run
            rPr = OxmlElement('w:rPr')
            
            # Add color
            color = OxmlElement('w:color')
            color.set(qn('w:val'), '0000FF')  # Blue color
            rPr.append(color)
            
            # Add underline
            underline = OxmlElement('w:u')
            underline.set(qn('w:val'), 'single')
            rPr.append(underline)
            
            run.append(rPr)
            
            # Add text element
            t = OxmlElement('w:t')
            t.text = text
            run.append(t)
            
            hyperlink.append(run)
            paragraph._p.append(hyperlink)
            
            logger.debug(f"Added hyperlink: {text} -> {url}")
            
        except Exception as e:
            logger.error(f"Error adding hyperlink: {str(e)}")
            # Fallback to plain text if hyperlink fails
            paragraph.add_run(text)
    
    def _apply_styles(
        self,
        element: BeautifulSoup,
        target: Union[Paragraph, Run]
    ) -> None:
        """Apply HTML styles to paragraph or run."""
        # Only process elements with 'style' attribute
        style = element.get('style')
        if not style or not isinstance(style, str):
            return
        
        for declaration in style.split(';'):
            if ':' not in declaration:
                continue
            
            prop, value = declaration.split(':')
            prop = prop.strip().lower()
            value = value.strip().lower()
            
            if prop == 'color' and isinstance(target, Run):
                if m := re.match(r'^#?([0-9a-f]{6})$', value):
                    hex_color = m.group(1)
                    r, g, b = tuple(int(hex_color[i:i+2], 16) for i in (0, 2, 4))
                    target.font.color.rgb = RGBColor(r, g, b)
            elif prop == 'text-align' and isinstance(target, Paragraph):
                alignments = {
                    'left': WD_ALIGN_PARAGRAPH.LEFT,
                    'center': WD_ALIGN_PARAGRAPH.CENTER,
                    'right': WD_ALIGN_PARAGRAPH.RIGHT,
                    'justify': WD_ALIGN_PARAGRAPH.JUSTIFY
                }
                if value in alignments:
                    target.alignment = alignments[value]
    
    _element_handlers = {
        'p': _handle_paragraph,
        'h1': _handle_heading,
        'h2': _handle_heading,
        'h3': _handle_heading,
        'h4': _handle_heading,
        'h5': _handle_heading,
        'h6': _handle_heading,
        'ul': lambda self, e: self._handle_list(e, False),
        'ol': lambda self, e: self._handle_list(e, True),
        'table': _handle_table,
        'img': _handle_image
    }